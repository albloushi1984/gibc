import streamlit as st
import pandas as pd
import os
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
from arabic_reshaper import reshape
from bidi.algorithm import get_display

APP_USER = "gibc"
APP_PASS = "ahmad@1984"

COMPANIES = [
    {
        "name": "الخليج العالمية لتشييد المباني",
        "key": "khalij",
        "employees_file": "موظفي_الخليج.csv",
        "logo": "logo_khalij.png",
        "sig_manager": "sig1.png"
    },
    {
        "name": "الشركة المصرية الكويتية لتشيد المباني",
        "key": "misrkwt",
        "employees_file": "موظفي_المصرية_الكويتية.csv",
        "logo": "logo_misr_kwt.png",
        "sig_manager": "sig1_misrkwt.png"
    },
    {
        "name": "شركة الذكاء العالي لاصلاح وصيانه الحواسيب الشخصيه او المحموله",
        "key": "ai",
        "employees_file": "موظفي_الذكاء.csv",
        "logo": "logo_ai.png",
        "sig_manager": "sig1_ai.png"
    },
    {
        "name": "شركة جلف هوم العالميه للاستيراد والتصدير",
        "key": "gulfhome",
        "employees_file": "موظفي_جلفهوم.csv",
        "logo": "logo_gulfhome.png",
        "sig_manager": "sig1_gulfhome.png"
    },
    {
        "name": "شركة هارموني كيدز لبيع ملابس الاطفال",
        "key": "harmonykids",
        "employees_file": "موظفي_هارمونيكيدز.csv",
        "logo": "logo_harmonykids.png",
        "sig_manager": "sig1_harmonykids.png"
    },
    {
        "name": "شركة جلف العالميه العقاريه",
        "key": "realestate",
        "employees_file": "موظفي_جلف_العقارية.csv",
        "logo": "logo_realestate.png",
        "sig_manager": "sig1_realestate.png"
    }
]
MANAGERS = ["أحمد محمد عباس البلوشي", "منة الله احمد محمود السيد"]
SIG_AHMED = "sig2.png"
EMP_SIG = "employee_sig.png"
FONT_PATH = "arial.ttf"
OUTPUT_PDF = "خطابات_الموظفين"
OUTPUT_WORD = "عقود_العمال_word"
EMP_COLS = ["الاسم", "الرقم المدني", "الجنسية", "المهنة"]
LABOR_OFFICES = [
    "إدارة عمل حولى", "إدارة عمل العاصمة", "إدارة عمل الفروانية"
]

os.makedirs(OUTPUT_PDF, exist_ok=True)
os.makedirs(OUTPUT_WORD, exist_ok=True)

def fix_arabic(text):
    return get_display(reshape(str(text).strip()))

def load_employees(filename):
    if not os.path.isfile(filename):
        return pd.DataFrame(columns=EMP_COLS)
    try:
        df = pd.read_csv(filename)
        for col in EMP_COLS:
            if col not in df.columns:
                df[col] = ""
        return df[EMP_COLS]
    except Exception:
        return pd.DataFrame(columns=EMP_COLS)

def save_employees(filename, df):
    df.to_csv(filename, index=False, encoding="utf-8")

def sanitize_filename(s):
    invalid = '<>:"/\\|?*'
    for ch in invalid:
        s = s.replace(ch, '_')
    return s.strip().replace(' ', '_')

def get_signature_path(company, manager):
    for c in COMPANIES:
        if c["name"] == company and manager == "منة الله احمد محمود السيد":
            return c.get("sig_manager", SIG_AHMED)
    return SIG_AHMED

class PDF(FPDF):
    def __init__(self, logo_file='', company_name='', *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.logo_file = logo_file
        self.company_name = company_name
        self.add_font("Arial", "", FONT_PATH, uni=True)
        self.add_font("Arial", "B", FONT_PATH, uni=True)
    def header(self):
        if self.logo_file and os.path.exists(self.logo_file):
            self.image(self.logo_file, x=150, y=8, w=40)
        self.set_font("Arial", "B", 14)
        self.cell(0, 10, fix_arabic(self.company_name), align="R", ln=1)
        self.ln(10)

def make_letter_pdf(comp, name, cid, job, salary, start, end, reason, mode, manager, manual_date, deduction, deduction_reason):
    pdf = PDF(logo_file=comp["logo"], company_name=comp["name"], orientation='P', unit='mm', format='A4')
    pdf.set_font("Arial", "", 18)
    pdf.add_page()
    pdf.ln(20)
    if mode == "شهادة راتب":
        pdf.cell(0, 14, fix_arabic("شهادة راتب"), ln=True, align="C")
        pdf.ln(8)
        paragraph = (
            f"تشهد {comp['name']} بأن الموظف يعمل: {name} / {cid}\n"
            f"يعمل لدينا بوظيفة: {job}\n"
            f"ويتقاضى راتبًا شهريًا قدره: {salary}\n"
            f"وذلك بناءً على طلبه لتقديمها إلى من يهمه الأمر."
        )
        pdf.multi_cell(0, 8, fix_arabic(paragraph), align="C")
        if deduction and deduction_reason:
            pdf.ln(4)
            line = f"تم خصم مبلغ قدره {deduction} من الموظف بسبب: {deduction_reason}"
            pdf.multi_cell(0, 8, fix_arabic(line), align="C")
    elif mode == "إجازة بدون راتب":
        pdf.cell(0, 14, fix_arabic("إجازة بدون راتب"), ln=True, align="C")
        pdf.ln(8)
        full_text = (
            f"نفيدكم بأن الموظف: {name} / {cid}\n"
            f"قد تقدم بطلب إجازة بدون راتب للفترة من {start} إلى {end}.\n"
            f"نظرًا لتمتع الموظف بإجازة خلال الفترة المحددة، لم يتم تحويل راتبه عن تلك الفترة."
        )
        pdf.multi_cell(0, 8, fix_arabic(full_text), align="C")
        if deduction and deduction_reason:
            pdf.ln(4)
            line = f"تم خصم مبلغ قدره {deduction} من الموظف بسبب: {deduction_reason}"
            pdf.multi_cell(0, 8, fix_arabic(line), align="C")
    elif mode == "خصم":
        pdf.cell(0, 14, fix_arabic("خطاب خصم"), ln=True, align="C")
        pdf.ln(8)
        full_text = (
            f"الموضوع: خصم من راتب موظف\n\n"
            f"نحيطكم علمًا بأنه قد تقرر خصم مبلغ وقدره ({deduction}) دينار كويتي من راتب الموظف:\n"
            f"{name} / {cid}\n"
            f"وذلك بتاريخ: {manual_date}\n"
            f"بسبب: {deduction_reason}.\n\n"
            f"يرجى التكرم باتخاذ اللازم وتنفيذ الخصم في كشف رواتب الشهر الحالي.\n\n"
            f"وتفضلوا بقبول فائق الاحترام والتقدير،،،"
        )
        pdf.multi_cell(0, 8, fix_arabic(full_text), align="C")
    else:  # إجازة عادية
        pdf.cell(0, 14, fix_arabic("إجازة"), ln=True, align="C")
        pdf.ln(8)
        full_text = (
            f"نفيدكم بأن الموظف: {name} / {cid}\n"
            f"قد قدم على إجازة خلال الفترة من {start} إلى {end}."
        )
        pdf.multi_cell(0, 8, fix_arabic(full_text), align="C")
        if deduction and deduction_reason:
            pdf.ln(4)
            line = f"تم خصم مبلغ قدره {deduction} من الموظف بسبب: {deduction_reason}"
            pdf.multi_cell(0, 8, fix_arabic(line), align="C")
    pdf_bytes = pdf.output(dest='S').encode('latin1')
    file_stream = BytesIO(pdf_bytes)
    return file_stream

def generate_contract_docx(data):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    style.font.size = Pt(13)
    for style_name in ['Normal', 'Table Grid']:
        if style_name in doc.styles:
            doc.styles[style_name].font.name = 'Arial'
            doc.styles[style_name]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    def align_rtl(p):
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.runs[0] if p.runs else p.add_run("")
        run.font.rtl = True
    def add_bond(title, content):
        p = doc.add_paragraph(f"{title}\n{content}")
        align_rtl(p)
    # --- إضافة شعار الشركة أعلى الصفحة ---
    comp = next((c for c in COMPANIES if c["name"] == data["company"]), None)
    if comp and os.path.exists(comp["logo"]):
        header = doc.sections[0].header
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(comp["logo"], width=Inches(1.4))
        paragraph.alignment = 1  # وسط

    p = doc.add_paragraph("نموذج عقد عمل في القطاع الأهلي"); p.alignment = 1
    p = doc.add_paragraph("دولة الكويت"); p.alignment = 1
    p = doc.add_paragraph(f"الهيئة العامة للقوى العاملة / {data['labor_office']}"); p.alignment = 1
    p = doc.add_paragraph(f"إنه في يوم {data['contract_date']}"); p.alignment = 1
    p = doc.add_paragraph("تحرر هذا العقد بين كل من :"); align_rtl(p)
    p = doc.add_paragraph(
        f"1- {data['company']}\nويمثلها في التوقيع على العقد:\nالاسم: {data['manager']}\n"
        f"رقم مدني / {data['manager_civilid']}\n             \" طرف اول \""
    ); align_rtl(p)
    p = doc.add_paragraph(
        f"2- الاسم: {data['worker_name']}\n"
        f"الجنسية: {data['nationality']}\n"
        f"رقم مدني: {data['worker_civilid']}\n"
        f"رقم الجواز: {data['passport']}\n"
        f"             \" طرف ثان \""
    ); align_rtl(p)
    p = doc.add_paragraph(
        "تمهيد\n"
        f"يمتلك الطرف الأول منشأة باسم/ {data['company']} تعمل في مجال المباني "
        f"ويرغب فى التعاقد مع الطرف الثاني للعمل لديه بمهنة {data['job']} وبعد أن أقر الطرفان بأهليتهما في إبرام هذا العقد تم الاتفاق علي ما يلي:"
    ); align_rtl(p)
    add_bond("البند الأول", "يعتبر التمهيد السابق جزءا لا يتجزأ من هذا العقد .")
    add_bond("البند الثاني", f"\" طبيعة العمل \"\nتعاقد الطرف الأول مع الطرف الثاني للعمل لديه بمهنة {data['job']} داخل دولة الكويت.")
    add_bond("البند الثالث", "\" فترة التجربة \"\n يخضع الطرف الثاني لفترة تجربة لمدة لا تزيد عن 100 يوم عمل ، ويحق لكل طرف إنهاء العقد خلال تلك الفترة دون إخطار ")
    add_bond("البند الرابع", f"\" قيمة الأجر \"\nيتقاضى الطرف الثاني عن تنفيذ هذا العقد أجرا مقداره {data['salary']} دينارا يدفع في نهاية كل شهر ولا يجوز للطرف الأول تخفيض الأجر أثناء سريان هذا العقد . ولا يجوز نقل الطرف الثاني إلى الأجر اليومي دون موافقته . ")
    doc.add_page_break()
    add_bond("البند الخامس", f"\" نفاذ العقد \"\نبدأ نفاذ العقد اعتبارا من {data['start_date']} ويلتزم الطرف الثاني بالقيام بأداء عمله طوال مدة نفاذة")
    add_bond("البند السادس", f"\" مدة العقد \"\n-هذا العقد غير محدد المدة ويبدأ اعتبارا من {data['real_start']} ولمدة  سنة  ، ويجوز تجديد العقد بموافقة الطرفين لمدد مماثلة بحد أقصى سنة  ميلادية.")
    add_bond("البند السابع", "\" الإجازة السنوية  \"\نللطرف الثاني الحق في إجازة سنوية مدفوعة الأجر مدتها 30 يوما ، ولا يستحقها عن السنة الأولى إلا بعد انقضاء مدة تسعة أشهر تحسب من تاريخ نفاذ العقد .")
    add_bond("البند الثامن", "\" عدد ساعات العمل \"\نلا يجوز للطرف الأول تشغيل الطرف الثاني لمدة تزيد عن ثماني ساعات عمل يوميا تتخللها فترة راحة لا تقل عن ساعة باستثناء الحالات المقررة قانونا .")
    add_bond("البند التاسع", "\" قيمة تذكرة السفر \"\نيتحمل الطرف الأول مصاريف عودة الطرف الثاني إلى بلده عند انتهاء علاقة العمل ومغادرته نهائيا للبلاد.")
    add_bond("البند العاشر", "\" التأمين ضد إصابات وأمراض العمل \"\نيلتزم الطرف الأول بالتأمين على الطرف الثانى ضد إصابات وأمراض العمل ، كما يلتزم بقيمة التأمين الصحى طبقا للقانون رقم (1) لسنة 1999  . ")
    add_bond("البند الحادى عشر", "\" مكافأة نهاية الخدمة \"\نيستحق الطرف الثان مكافأة نهاية الخدمة المنصوص عليها بالقوانين المنظمة ")
    add_bond("البند الثانى عشر", "\" القانون الواجب التطبيق \"\ن تسري أحكام قانون العمل في القطاع الأهلي رقم 6 لسنة 2010 والقرارات المنفذة له فيما لم يرد بشأنه نص في هذا العقد ، ويقع باطلا كل شرط تم الاتفاق عليه بالمخالفة لأحكام القانون ، ما لم يكن فيه ميزة أفضل للعامل .")
    add_bond("البند الثالث عشر", "\"شروط خاصة \"\ن1 لا  \n2 لا     \n3 لا  ")
    add_bond("البند الرابع عشر", "\" المحكمة المختصة \"\نختص المحكمة الكلية ودوائرها العمالية طبقا لأحكام القانون رقم 46 لسنة 1987 ، بنظر كافة المنازعات الناشئة عن تطبيق أو تفسير هذا العقد.")
    add_bond("البند الخامس عشر", "\" لغة العقد \"\نحرر هذا العقد باللغه  العربيه  ، ويعتد بنصوص اللغة العربية عند وقوع أى تعارض .")
    add_bond("البند السادس عشر", "\" نسخ العقد \"\نحرر هذا العقد من ثلاث نسخ بيد كل طرف نسخة للعمل بموجبها والثالثة تودع لدى الهيئة العامة للقوى العاملة.")
    doc.add_paragraph("\n\n")
    table = doc.add_table(rows=1, cols=2)
    row = table.rows[0]
    cell1, cell2 = row.cells
    cell1.paragraphs[0].add_run("الطرف الأول").bold = True
    cell2.paragraphs[0].add_run("الطرف الثاني").bold = True
    cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    sign_path = data['manager_sign']
    if os.path.exists(sign_path):
        try:
            run = cell1.paragraphs[0].add_run()
            run.add_picture(sign_path, width=Inches(0.8))
        except Exception:
            pass
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def login_form():
    st.markdown("""
    <div style='text-align:center; color:#1976d2; font-family:"Cairo",Arial,sans-serif'>
        <img src="https://cdn-icons-png.flaticon.com/512/3135/3135715.png" width="80"><br>
        <h2>بوابة إدارة الموظفين والعقود</h2>
    </div>
    """, unsafe_allow_html=True)
    with st.form("login_form"):
        user = st.text_input("اسم المستخدم")
        passwd = st.text_input("كلمة المرور", type="password")
        submitted = st.form_submit_button("دخول")
        if submitted:
            if user == APP_USER and passwd == APP_PASS:
                st.session_state["logged_in"] = True
                st.success("تم تسجيل الدخول بنجاح!")
                st.experimental_rerun()
            else:
                st.error("اسم المستخدم أو كلمة المرور غير صحيحة")
    st.stop()

if "logged_in" not in st.session_state:
    st.session_state["logged_in"] = False

if not st.session_state["logged_in"]:
    login_form()

st.set_page_config(page_title="نظام إدارة الموظفين والعقود", layout="wide", page_icon=":briefcase:")
st.markdown(
    """
    <style>
    html, body, [class*="css"] {
        font-family: 'Cairo', 'Arial', sans-serif !important;
        direction: rtl;
    }
    .stTabs [data-baseweb="tab"] {font-size:20px; font-family:"Cairo"; direction:rtl}
    .stApp {background-color:#f6fafd;}
    .stButton>button {
        background: linear-gradient(90deg,#1976d2,#43a047);
        color:white;
        font-size:18px;
        border-radius:8px;
        padding: 0.5em 2em;
        margin-bottom:3px;
        margin-top:7px;
        font-family: 'Cairo', 'Arial', sans-serif;
        transition:0.15s;
    }
    .stButton>button:hover {
        background:linear-gradient(90deg,#43a047,#1976d2);
        color:#fff;
    }
    .stTextInput>div>input, .stSelectbox>div>div {
        font-size: 17px;
        background: #e3eafc;
        border-radius:6px;
        color:#17408b;
    }
    .title-main {
        color:#1976d2;font-size:2.3rem;font-weight:900;margin-bottom:10px;text-align:center;font-family:'Cairo',Arial,sans-serif
    }
    .stAlert {
        background:#e3fcec !important;
        color:#257a3e !important;
        font-size:18px;
    }
    </style>
    <link href="https://fonts.googleapis.com/css?family=Cairo&display=swap" rel="stylesheet">
    """, unsafe_allow_html=True
)

st.markdown('<div class="title-main">نظام إدارة الموظفين والعقود والخطابات</div>', unsafe_allow_html=True)

# --- شعار الشركة أعلى الصفحة حسب التبويب الحالي ---
company_logo = None
if "company_logo" not in st.session_state:
    st.session_state["company_logo"] = None

tabs = st.tabs(["👥 إدارة الموظفين", "📄 خطابات وإجازات", "📝 توليد عقد عمل"🏢 شهادات الشركات"])

with tabs[0]:
    st.markdown("<h3 style='color:#1976d2'>إدارة جميع الموظفين</h3>", unsafe_allow_html=True)
    company_names = [c["name"] for c in COMPANIES]
    comp_idx = st.selectbox("اختر الشركة:", range(len(company_names)), format_func=lambda i: company_names[i])
    comp = COMPANIES[comp_idx]
    # شعار الشركة في الأعلى
    if os.path.exists(comp["logo"]):
        st.image(comp["logo"], width=110)
        st.session_state["company_logo"] = comp["logo"]
    df = load_employees(comp["employees_file"])
    st.dataframe(df, use_container_width=True, hide_index=True)

    with st.expander("➕ إضافة موظف جديد"):
        with st.form("add_emp_form"):
            c1, c2 = st.columns(2)
            name = c1.text_input("الاسم")
            cid = c2.text_input("الرقم المدني")
            nationality = c1.text_input("الجنسية")
            job = c2.text_input("المهنة")
            submitted = st.form_submit_button("إضافة")
            if submitted:
                if name and cid:
                    new_row = {"الاسم": name, "الرقم المدني": cid, "الجنسية": nationality, "المهنة": job}
                    df = pd.concat([df, pd.DataFrame([new_row])], ignore_index=True)
                    save_employees(comp["employees_file"], df)
                    st.success("تمت إضافة الموظف!")
                    st.experimental_rerun()
                else:
                    st.error("يرجى إدخال الاسم والرقم المدني.")

    with st.expander("🗑️ حذف موظف"):
        emp_names = df["الاسم"].tolist()
        emp_to_delete = st.selectbox("اختر الموظف للحذف:", [""] + emp_names)
        if st.button("حذف الموظف", key="del_emp_btn") and emp_to_delete:
            df = df[df["الاسم"] != emp_to_delete]
            save_employees(comp["employees_file"], df)
            st.success("تم حذف الموظف.")
            st.experimental_rerun()

with tabs[1]:
    st.markdown("<h3 style='color:#1976d2'>توليد خطابات وإجازات</h3>", unsafe_allow_html=True)
    company_names = [c["name"] for c in COMPANIES]
    comp_idx2 = st.selectbox("اختر الشركة:", range(len(company_names)), format_func=lambda i: company_names[i], key="pdf_company")
    comp2 = COMPANIES[comp_idx2]
    if os.path.exists(comp2["logo"]):
        st.image(comp2["logo"], width=110)
        st.session_state["company_logo"] = comp2["logo"]
    df2 = load_employees(comp2["employees_file"])
    emp_list = df2["الاسم"].tolist()
    emp_name = st.selectbox("اسم الموظف", [""] + emp_list)
    emp_cid = st.text_input("الرقم المدني", value=df2[df2["الاسم"]==emp_name]["الرقم المدني"].iloc[0] if emp_name else "")
    job = st.text_input("المسمى الوظيفي", value=df2[df2["الاسم"]==emp_name]["المهنة"].iloc[0] if emp_name else "")
    salary = st.text_input("الراتب الشهري")
    start = st.text_input("تاريخ البداية (YYYY-MM-DD)")
    end = st.text_input("تاريخ النهاية (YYYY-MM-DD)")
    reason = st.text_input("سبب الإجازة")
    manager = st.selectbox("اختر المدير:", MANAGERS)
    manual_date = st.text_input("تاريخ الخطاب (يدوي)")
    deduction = st.text_input("مبلغ الخصم (إذا وجد)")
    deduction_reason = st.text_input("سبب الخصم (إذا وجد)")
    letter_type = st.selectbox("نوع الخطاب", ["إجازة", "إجازة بدون راتب", "شهادة راتب", "خصم"])
    if st.button("توليد الخطاب"):
        if emp_name and emp_cid:
            pdf_file = make_letter_pdf(
                comp2, emp_name, emp_cid, job, salary, start, end, reason, letter_type, manager, manual_date, deduction, deduction_reason
            )
            fname = f"{letter_type}_{sanitize_filename(emp_name)}.pdf"
            st.success("✅ تم توليد الخطاب بنجاح!")
            st.download_button("تحميل الخطاب PDF", pdf_file, file_name=fname, use_container_width=True)
        else:
            st.error("يرجى اختيار الموظف وملء البيانات.")

with tabs[2]:
    st.markdown("<h3 style='color:#1976d2'>توليد عقد عمل (Word)</h3>", unsafe_allow_html=True)
    company_names = [c["name"] for c in COMPANIES]
    col1, col2 = st.columns(2)
    with col1:
        labor_office = st.selectbox("إدارة العمل:", LABOR_OFFICES)
        company = st.selectbox("اسم الشركة", company_names, key="contract_company")
        manager = st.selectbox("اسم المدير", MANAGERS)
        manager_civilid = st.text_input("الرقم المدني للمدير")
        worker_name = st.text_input("اسم العامل")
        nationality = st.text_input("الجنسية")
        job = st.text_input("المهنة")
        comp3 = next((c for c in COMPANIES if c["name"] == company), None)
        if comp3 and os.path.exists(comp3["logo"]):
            st.image(comp3["logo"], width=110)
            st.session_state["company_logo"] = comp3["logo"]
    with col2:
        worker_civilid = st.text_input("الرقم المدني للعامل")
        passport = st.text_input("رقم الجواز")
        salary = st.text_input("الراتب")
        contract_date = st.text_input("تاريخ العقد (YYYY-MM-DD)")
        start_date = st.text_input("تاريخ بداية النفاذ (YYYY-MM-DD)")
        real_start = st.text_input("تاريخ بداية العقد (YYYY-MM-DD)")
    if st.button("توليد العقد"):
        if all([labor_office, company, manager, manager_civilid, worker_name, nationality, worker_civilid, passport, job, salary, contract_date, start_date, real_start]):
            data = {
                "labor_office": labor_office, "company": company, "manager": manager,
                "manager_civilid": manager_civilid, "worker_name": worker_name, "worker_civilid": worker_civilid,
                "nationality": nationality, "passport": passport, "job": job, "salary": salary,
                "contract_date": contract_date, "start_date": start_date, "real_start": real_start,
                "manager_sign": get_signature_path(company, manager)
            }
            docx_file = generate_contract_docx(data)
            file_name = f"عقد_{sanitize_filename(company)}_{sanitize_filename(worker_name)}_{sanitize_filename(worker_civilid)}.docx"
            st.success("✅ تم توليد العقد بنجاح!")
            st.download_button("تحميل ملف العقد", docx_file, file_name=file_name, use_container_width=True)
        else:
            st.error("يرجى إدخال جميع البيانات.")

with tabs[3]:
    st.markdown("<h3 style='color:#1976d2'>جميع الشركات - تحميل الرخصة واعتماد التوقيع</h3>", unsafe_allow_html=True)

    companies_cert = [
        {
            "name": "شركه الخليج العالميه لتشيد المباني",
            "license": "رخصه الخليج.pdf",
            "attestation": "اعتماد الخليج.pdf"
        },
        {
            "name": "الشركه المصريه الكويتيه لتشيد المباني",
            "license": "رخصه المصريه.pdf",
            "attestation": "اعتماد المصريه الكويتيه.pdf"
        },
        {
            "name": "الذكاء العالي لاصلاح وصيانه الحواسيب الشخصيه والمحموله",
            "license": "رخصه الذكاء.pdf",
            "attestation": "اعتماد الذكاء.pdf"
        }
        # أضف شركات أخرى هنا بنفس الأسلوب إذا احتجت لاحقًا
    ]

    for comp in companies_cert:
        st.markdown(f"### {comp['name']}")
        col1, col2 = st.columns(2)
        with col1:
            try:
                with open(comp["license"], "rb") as f:
                    st.download_button(
                        label="تحميل الرخصة",
                        data=f,
                        file_name=comp["license"],
                        mime="application/pdf"
                    )
            except FileNotFoundError:
                st.warning("ملف الرخصة غير موجود.")
        with col2:
            try:
                with open(comp["attestation"], "rb") as f:
                    st.download_button(
                        label="تحميل اعتماد التوقيع",
                        data=f,
                        file_name=comp["attestation"],
                        mime="application/pdf"
                    )
            except FileNotFoundError:
                st.warning("ملف الاعتماد غير موجود.")
        st.markdown("---")